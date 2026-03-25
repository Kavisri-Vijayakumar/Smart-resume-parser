<<<<<<< HEAD
import streamlit as st
=======
>>>>>>> 647fb528e9c4a3e8ebab0d8c18c3e684bb45c6ec
import pdfplumber
import docx
import re
import json
<<<<<<< HEAD
import tempfile


# ===== Extract Text =====
def extract_text(file):
    text = ""

    if file.name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + "\n"
=======


def extract_text(file_path):
    text = ""

    if file_path.endswith(".pdf"):
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            return f"Error reading PDF: {e}"

    elif file_path.endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            return f"Error reading DOCX: {e}"

    else:
        return "Unsupported file format"
>>>>>>> 647fb528e9c4a3e8ebab0d8c18c3e684bb45c6ec

    return text


<<<<<<< HEAD
# ===== Clean Text =====
=======
>>>>>>> 647fb528e9c4a3e8ebab0d8c18c3e684bb45c6ec
def clean_text(text):
    text = re.sub(r'\n+', '\n', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


<<<<<<< HEAD
# ===== Extract Skills =====
=======
>>>>>>> 647fb528e9c4a3e8ebab0d8c18c3e684bb45c6ec
def extract_skills(text):
    skills_list = [
        "Python", "Java", "C++", "SQL",
        "Machine Learning", "HTML", "CSS",
        "JavaScript", "Data Science"
    ]

<<<<<<< HEAD
    found_skills = []

    for skill in skills_list:
        if skill.lower() in text.lower():
            found_skills.append(skill)

    return found_skills


# ===== Extract Education =====
def extract_education(text):
    keywords = ["Bachelor", "B.Tech", "BSc", "BCA", "Master", "MBA"]

    found = []
    for word in keywords:
        if word.lower() in text.lower():
            found.append(word)

    return list(set(found))


# ===== Extract Experience =====
def extract_experience(text):
    keywords = ["Intern", "Experience", "Project", "Worked"]

    found = []
    for word in keywords:
        if word.lower() in text.lower():
            found.append(word)

    return list(set(found))


# ===== UI =====
st.title("📄 Smart Resume Parser")

uploaded_file = st.file_uploader("Upload your Resume", type=["pdf", "docx"])

if uploaded_file is not None:
    text = extract_text(uploaded_file)
    cleaned = clean_text(text)

    skills = extract_skills(cleaned)
    education = extract_education(cleaned)
    experience = extract_experience(cleaned)

    result = {
        "Skills": skills,
        "Education": education,
        "Experience": experience
    }

    st.subheader("📌 Extracted Information")
    st.json(result)

    st.download_button(
        "Download Result",
        json.dumps(result, indent=4),
        file_name="result.json"
    )
=======
    return [skill for skill in skills_list if skill.lower() in text.lower()]


def extract_education(text):
    keywords = [
        "Bachelor", "B.Tech", "BSc", "BCA",
        "Master", "M.Tech", "MSc", "MBA"
    ]

    return list(set([k for k in keywords if k.lower() in text.lower()]))


def extract_experience(text):
    keywords = ["Intern", "Experience", "Worked", "Project"]

    return list(set([k for k in keywords if k.lower() in text.lower()]))


if __name__ == "__main__":
    file_path = "sample.pdf"

    text = extract_text(file_path)

    if "Error" in text or "Unsupported" in text:
        print(text)
    else:
        cleaned = clean_text(text)

        result = {
            "Skills": extract_skills(cleaned),
            "Education": extract_education(cleaned),
            "Experience": extract_experience(cleaned)
        }

        print("\n===== FINAL OUTPUT =====\n")
        print(json.dumps(result, indent=4))
>>>>>>> 647fb528e9c4a3e8ebab0d8c18c3e684bb45c6ec
